import re
import io
import logging
from typing import Dict, List, Tuple, Any, Optional
import fitz  # PyMuPDF
import docx
from backend.app.schemas.models import (
    ExtractedResume, ProjectItem, ExperienceItem, 
    EducationItem, CertificationItem, ResumeScoreBreakdown
)

logger = logging.getLogger("parser")

# Standard taxonomy of common technical skills
KNOWN_SKILLS = {
    # Languages
    "python", "javascript", "typescript", "java", "c++", "c#", "c", "go", "golang", "rust", 
    "ruby", "php", "swift", "kotlin", "scala", "r", "dart", "sql", "html", "html5", "css", "css3",
    # Frontend
    "react", "react.js", "reactjs", "next.js", "nextjs", "vue", "vue.js", "vuejs", "angular", 
    "svelte", "tailwind", "tailwind css", "bootstrap", "redux", "zustand", "graphql", "webpack", "vite",
    # Backend
    "node.js", "nodejs", "express", "express.js", "fastapi", "flask", "django", "spring", 
    "spring boot", "nestjs", "asp.net", ".net", "laravel", "rails", "rest api", "restful apis", "grpc",
    # Databases
    "mongodb", "postgresql", "postgres", "mysql", "sqlite", "redis", "cassandra", "dynamodb", 
    "mariadb", "oracle", "neo4j", "elasticsearch", "firebase", "supabase", "firestore",
    # Cloud & DevOps
    "aws", "amazon web services", "azure", "gcp", "google cloud", "docker", "kubernetes", "k8s", 
    "ci/cd", "github actions", "jenkins", "terraform", "ansible", "nginx", "linux", "git", "github", "gitlab",
    # Computer Vision & AI/ML
    "yolo", "yolov8", "yolov5", "opencv", "cv2", "machine learning", "deep learning", "nlp", "computer vision", 
    "tensorflow", "pytorch", "keras", "scikit-learn", "pandas", "numpy", "power bi", "tableau", "spark", 
    "hadoop", "kafka", "llm", "langchain", "transformers", "huggingface", "data analysis", "data engineering",
    "onnx", "tensorrt", "cnn", "resnet", "xgboost", "matplotlib", "seaborn", "nltk", "spacy",
    # Concepts & Architecture
    "microservices", "system design", "distributed systems", "oop", "object oriented programming",
    "data structures", "algorithms", "agile", "scrum", "unit testing", "tdd", "ci/cd pipelines"
}

CATEGORY_MAP = {
    "Languages": ["python", "javascript", "typescript", "java", "c++", "c#", "c", "go", "golang", "rust", "ruby", "php", "swift", "kotlin", "sql", "html", "css", "dart", "r"],
    "Frontend": ["react", "react.js", "reactjs", "next.js", "nextjs", "vue", "vue.js", "angular", "tailwind", "tailwind css", "bootstrap", "redux", "vite", "webpack"],
    "Backend & APIs": ["node.js", "nodejs", "fastapi", "flask", "django", "express", "spring boot", "rest api", "graphql", "grpc", "microservices"],
    "Databases": ["mongodb", "postgresql", "postgres", "mysql", "redis", "sqlite", "elasticsearch", "dynamodb", "supabase", "firebase"],
    "Cloud & DevOps": ["aws", "azure", "gcp", "docker", "kubernetes", "git", "github", "ci/cd", "linux", "terraform", "jenkins"],
    "AI & Data Science": ["yolo", "yolov8", "opencv", "machine learning", "deep learning", "nlp", "computer vision", "pandas", "numpy", "scikit-learn", "pytorch", "tensorflow", "power bi", "tableau", "llm", "xgboost"]
}

class ResumeParser:
    @staticmethod
    def extract_text_from_pdf(pdf_bytes: bytes) -> str:
        text = ""
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            for page in doc:
                text += page.get_text() + "\n"
        except Exception as e:
            logger.error(f"Failed to extract text with PyMuPDF: {e}")
        return text.strip()

    @staticmethod
    def extract_text_from_docx(docx_bytes: bytes) -> str:
        text = ""
        try:
            doc = docx.Document(io.BytesIO(docx_bytes))
            for p in doc.paragraphs:
                if p.text:
                    text += p.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            logger.error(f"Failed to extract text from docx: {e}")
        return text.strip()

    @staticmethod
    def clean_text(text: str) -> str:
        text = re.sub(r'\r\n|\r', '\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    @staticmethod
    def extract_contact_info(text: str) -> Dict[str, Optional[str]]:
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Name heuristic: typically first non-empty line with 2-4 alphabetic words
        name = "Candidate"
        for line in lines[:5]:
            clean_line = re.sub(r'[^a-zA-Z\s]', '', line).strip()
            words = clean_line.split()
            if 2 <= len(words) <= 4 and not any(kw in clean_line.lower() for kw in ["resume", "curriculum", "profile", "contact", "email", "phone", "developer", "engineer"]):
                name = clean_line.title()
                break

        # Email
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
        email = email_match.group(0) if email_match else None

        # Phone
        phone_match = re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
        phone = phone_match.group(0) if phone_match else None

        # Location heuristic
        location = None
        loc_patterns = [
            r'([A-Z][a-zA-Z\s]+,\s*[A-Z]{2}(\s*\d{5})?)',
            r'([A-Z][a-zA-Z\s]+,\s*(?:USA|India|Canada|UK|Germany|Australia))'
        ]
        for pat in loc_patterns:
            loc_match = re.search(pat, text)
            if loc_match:
                location = loc_match.group(1).strip()
                break

        return {"name": name, "email": email, "phone": phone, "location": location}

    @staticmethod
    def extract_skills(text: str) -> Tuple[List[str], Dict[str, List[str]]]:
        lower_text = " " + text.lower() + " "
        found_skills = set()

        for skill in KNOWN_SKILLS:
            # Word boundary regex for single words, direct substring for multi-word or symbols
            if len(skill.split()) == 1 and skill.isalpha():
                pattern = r'\b' + re.escape(skill) + r'\b'
                if re.search(pattern, lower_text):
                    found_skills.add(skill)
            else:
                if skill in lower_text:
                    found_skills.add(skill)

        # Standardize skill names (e.g. capitalize nicely)
        display_map = {
            "python": "Python", "javascript": "JavaScript", "typescript": "TypeScript", "java": "Java",
            "c++": "C++", "c#": "C#", "c": "C", "go": "Go", "golang": "Go", "rust": "Rust",
            "react": "React", "react.js": "React", "reactjs": "React", "next.js": "Next.js", "nextjs": "Next.js",
            "vue": "Vue.js", "angular": "Angular", "tailwind": "Tailwind CSS", "tailwind css": "Tailwind CSS",
            "node.js": "Node.js", "nodejs": "Node.js", "express": "Express.js", "fastapi": "FastAPI",
            "flask": "Flask", "django": "Django", "spring boot": "Spring Boot",
            "mongodb": "MongoDB", "postgresql": "PostgreSQL", "postgres": "PostgreSQL", "mysql": "MySQL",
            "sqlite": "SQLite", "redis": "Redis", "aws": "AWS", "docker": "Docker", "kubernetes": "Kubernetes",
            "git": "Git", "github": "GitHub", "machine learning": "Machine Learning", "deep learning": "Deep Learning",
            "nlp": "NLP", "pandas": "Pandas", "numpy": "NumPy", "scikit-learn": "Scikit-Learn",
            "power bi": "Power BI", "tableau": "Tableau", "rest api": "REST APIs", "graphql": "GraphQL",
            "microservices": "Microservices", "system design": "System Design", "ci/cd": "CI/CD",
            "html": "HTML5", "css": "CSS3", "sql": "SQL", "pytorch": "PyTorch", "tensorflow": "TensorFlow",
            "yolo": "YOLO", "yolov8": "YOLOv8", "yolov5": "YOLOv5", "opencv": "OpenCV", "cv2": "OpenCV",
            "onnx": "ONNX", "tensorrt": "TensorRT", "xgboost": "XGBoost", "keras": "Keras"
        }

        standardized_skills = []
        for s in found_skills:
            name = display_map.get(s, s.title())
            if name not in standardized_skills:
                standardized_skills.append(name)

        # Categorize
        categories = {}
        for cat, skills in CATEGORY_MAP.items():
            cat_list = []
            for s in skills:
                std_name = display_map.get(s, s.title())
                if std_name in standardized_skills and std_name not in cat_list:
                    cat_list.append(std_name)
            if cat_list:
                categories[cat] = cat_list

        return standardized_skills, categories

    @staticmethod
    def extract_sections(text: str) -> Dict[str, str]:
        section_headers = {
            "summary": ["summary", "professional summary", "about me", "objective", "profile", "career objective"],
            "experience": ["experience", "work experience", "employment history", "professional experience", "work history", "internships", "internship experience"],
            "projects": ["projects", "personal projects", "academic projects", "key projects", "notable projects", "technical projects", "project work", "project details", "selected projects", "major projects", "featured projects"],
            "skills": ["skills", "technical skills", "core competencies", "skills & tools", "technologies", "key skills", "skill set", "technical proficiencies", "technical expertise", "skills & abilities"],
            "education": ["education", "academic background", "academics", "qualifications", "educational background", "academic qualifications"],
            "certifications": ["certifications", "licenses", "certificates", "courses", "certifications & licenses", "online courses"],
            "achievements": ["achievements", "awards", "honors", "publications", "extracurricular", "achievements & awards", "honors & awards", "activities"]
        }

        # Sub-labels that should NEVER be treated as top-level section headers
        SUB_LABELS = {
            "skills", "technologies", "tools", "technologies used", "tools used", "tech stack",
            "description", "role", "key responsibilities", "responsibilities", "highlights", "duration",
            "environment", "frameworks", "libraries", "database", "databases", "frontend", "backend"
        }

        lines = text.split('\n')
        sections: Dict[str, List[str]] = {k: [] for k in section_headers}
        current_section = "summary"

        for line in lines:
            trimmed = line.strip()
            if not trimmed:
                continue

            lower_line = trimmed.lower().rstrip(':').strip()
            
            # Check if this line is a top-level section header
            detected = None
            if len(trimmed.split()) <= 4 and not trimmed.startswith(('•', '-', '*', '▪', '▫', '–', '—', '>', '●')) and lower_line not in SUB_LABELS:
                for sec, aliases in section_headers.items():
                    if lower_line in aliases:
                        detected = sec
                        break

            if detected:
                current_section = detected
            else:
                sections[current_section].append(trimmed)

        return {k: "\n".join(v).strip() for k, v in sections.items()}

    @classmethod
    def parse_projects(cls, projects_text: str, full_resume_text: str = "") -> List[ProjectItem]:
        if not projects_text and not full_resume_text:
            return []
        
        # Strategy 1: Parse the projects_text section
        items = cls._extract_projects_from_text(projects_text)

        # Strategy 2: If < 2 projects found, search the full resume text for project markers
        if len(items) < 2 and full_resume_text:
            full_scan_items = cls._extract_projects_from_text(full_resume_text)
            if len(full_scan_items) > len(items):
                items = full_scan_items

        return items[:6]

    @classmethod
    def _extract_projects_from_text(cls, text: str) -> List[ProjectItem]:
        if not text:
            return []

        lines = [l.strip() for l in text.split('\n') if l.strip()]
        if not lines:
            return []

        # Project candidate boundary detection
        project_blocks = []
        current_header = None
        current_lines = []

        def is_project_header(line: str, next_line: Optional[str] = None) -> bool:
            # Cannot start with standard bullet symbols
            if line.startswith(('•', '-', '*', '▪', '▫', '–', '—', '>', '●', '✦')):
                return False
            # Check for title patterns
            # 1. Has separators like | or – or - followed by tech/dates: "Project Name | React, Node.js"
            if re.search(r'\b[A-Za-z0-9\s]{3,50}\s*(?:\||–|—|-)\s*(?:[A-Za-z0-9,\s]{3,}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|20\d\d))', line):
                return True
            # 2. Numbered prefix: "1. Traffic Management" or "Project 1: ..." or "Project #1"
            if re.match(r'^(?:(?:\d+\.|\d+\)|\bProject\s*(?:#?\d+|[A-Z]):?))\s+[A-Za-z]', line, re.I):
                return True
            # 3. Contains parenthesis with date or tech: "Smart Traffic System (YOLOv8, OpenCV)"
            if re.search(r'^[A-Za-z0-9\s]{3,50}\s*\([A-Za-z0-9,\s\-\–]+\)', line):
                return True
            # 4. Short non-bullet line (< 60 chars) followed by bullet points
            if len(line.split()) <= 8 and len(line) < 65 and not line.endswith(('.', '!', '?')):
                if next_line and next_line.startswith(('•', '-', '*', '▪', '▫', '–', '—', '>', '●')):
                    return True
                if line.isupper() or (line.istitle() and len(line.split()) <= 6):
                    return True
            return False

        for i, line in enumerate(lines):
            next_l = lines[i + 1] if i + 1 < len(lines) else None
            if is_project_header(line, next_l):
                if current_header and current_lines:
                    project_blocks.append((current_header, current_lines))
                current_header = line
                current_lines = []
            else:
                if current_header:
                    current_lines.append(line)
                else:
                    if i == 0:
                        current_header = line
                    else:
                        current_lines.append(line)

        if current_header:
            project_blocks.append((current_header, current_lines))

        items = []
        for header, b_lines in project_blocks:
            clean_head = re.sub(r'^(?:\d+\.|\d+\)|\bProject\s*(?:#?\d+|[A-Z]):?)\s*', '', header, flags=re.I).strip()
            clean_head = clean_head.lstrip('•-*|–— ')
            parts = re.split(r'\||–|—|-|\(|\[', clean_head)
            title = parts[0].strip().rstrip(':')

            if len(title) < 3 or title.lower() in {"projects", "academic projects", "technical projects", "key highlights", "technologies used", "github", "description"}:
                continue

            full_block = header + "\n" + "\n".join(b_lines)
            skills, _ = cls.extract_skills(full_block)
            
            highlights = [l.lstrip('•-*▪▫–—>●✦ ') for l in b_lines if len(l.strip()) > 8]
            if not highlights and b_lines:
                highlights = [b_lines[0]]
            
            desc = " ".join(highlights[:2]) if highlights else f"Project focused on {', '.join(skills[:3]) if skills else title}."
            
            items.append(ProjectItem(
                title=title,
                description=desc[:300],
                technologies=skills,
                highlights=highlights[:5]
            ))

        return items

    @classmethod
    def parse_experience(cls, exp_text: str) -> List[ExperienceItem]:
        if not exp_text:
            return []
        
        items = []
        blocks = re.split(r'\n\s*\n', exp_text)
        
        for block in blocks:
            lines = [l.strip() for l in block.split('\n') if l.strip()]
            if not lines:
                continue
            
            first_line = lines[0].lstrip('•-* ')
            role = "Software Engineer"
            company = "Tech Organization"
            duration = None
            
            # Look for dates like 2022 - 2024, May 2021 - Present
            date_match = re.search(r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|20\d\d)\b.*?(?:Present|Current|20\d\d))', block, re.I)
            if date_match:
                duration = date_match.group(0)

            parts = re.split(r'\||–|-|,| at ', first_line)
            if len(parts) >= 2:
                role = parts[0].strip()
                company = parts[1].strip()
            elif len(parts) == 1:
                role = parts[0].strip()

            skills, _ = cls.extract_skills(block)
            bullets = [l.lstrip('•-*> ') for l in lines[1:] if len(l.strip()) > 10]

            if len(role) > 2:
                items.append(ExperienceItem(
                    role=role,
                    company=company,
                    duration=duration,
                    responsibilities=bullets[:5],
                    technologies=skills
                ))

        return items[:5]

    @classmethod
    def parse_education(cls, edu_text: str) -> List[EducationItem]:
        if not edu_text:
            return []
        
        items = []
        lines = [l.strip() for l in edu_text.split('\n') if l.strip()]
        for line in lines:
            if any(deg in line.lower() for deg in ["bachelor", "master", "b.tech", "b.e", "b.s", "m.s", "m.tech", "ph.d", "degree", "diploma", "computer science", "engineering"]):
                parts = re.split(r'\||–|-|,', line)
                degree = parts[0].strip()
                inst = parts[1].strip() if len(parts) > 1 else "University / Institution"
                year_match = re.search(r'\b(20\d\d|19\d\d)\b', line)
                year = year_match.group(0) if year_match else None
                items.append(EducationItem(degree=degree, institution=inst, year=year))
        
        if not items and edu_text:
            items.append(EducationItem(degree="Bachelor of Science in Computer Science", institution="University / College"))

        return items[:3]

    @classmethod
    def parse_resume(cls, text: str, filename: Optional[str] = None) -> ExtractedResume:
        cleaned = cls.clean_text(text)

        # 1. LLM-Powered parsing with Gemini (if API key is present)
        from backend.app.config import settings
        if settings.GEMINI_API_KEY:
            try:
                import json
                from google import genai
                client = genai.Client(api_key=settings.GEMINI_API_KEY)
                prompt = f"""You are an expert resume parsing engine.
Extract all structured data from the candidate's resume below with 100% accuracy.
CRITICAL INSTRUCTION: Extract ALL projects listed in the resume (typically 2 to 4 projects). Do NOT miss any project!

RESUME RAW TEXT:
{cleaned[:5000]}

OUTPUT VALID JSON ONLY with this schema:
{{
  "name": "Candidate Name",
  "email": "email or null",
  "phone": "phone or null",
  "location": "location or null",
  "summary": "1-2 sentence summary or null",
  "skills": ["Skill1", "Skill2", ...],
  "experience": [
    {{
      "role": "Role Title",
      "company": "Company Name",
      "duration": "Dates or null",
      "responsibilities": ["Bullet 1", "Bullet 2"],
      "technologies": ["Tech1", "Tech2"]
    }}
  ],
  "projects": [
    {{
      "title": "Exact Project Name",
      "description": "Brief summary",
      "technologies": ["Tech1", "Tech2"],
      "highlights": ["Bullet 1", "Bullet 2"]
    }}
  ],
  "education": [
    {{
      "degree": "Degree",
      "institution": "University",
      "year": "Year or null"
    }}
  ],
  "certifications": [{{"name": "Cert Name"}}],
  "achievements": ["Achievement 1"]
}}"""
                response = client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=prompt
                )
                raw_json = response.text.strip()
                if raw_json.startswith("```json"):
                    raw_json = raw_json.split("```json")[1].split("```")[0].strip()
                elif raw_json.startswith("```"):
                    raw_json = raw_json.split("```")[1].split("```")[0].strip()
                
                data = json.loads(raw_json)
                
                # Format into ExtractedResume
                extracted_skills, categories = cls.extract_skills(cleaned)
                # Merge any extra skills identified by LLM
                for s in data.get("skills", []):
                    if s not in extracted_skills:
                        extracted_skills.append(s)

                projects = [
                    ProjectItem(
                        title=p.get("title", "Project"),
                        description=p.get("description", "") or "Project implementation",
                        technologies=p.get("technologies", []),
                        highlights=p.get("highlights", [])
                    ) for p in data.get("projects", []) if p.get("title")
                ]

                experience = [
                    ExperienceItem(
                        role=e.get("role", "Software Engineer"),
                        company=e.get("company", "Company"),
                        duration=e.get("duration"),
                        responsibilities=e.get("responsibilities", []),
                        technologies=e.get("technologies", [])
                    ) for e in data.get("experience", []) if e.get("role")
                ]

                education = [
                    EducationItem(
                        degree=ed.get("degree", "Degree"),
                        institution=ed.get("institution", "University"),
                        year=ed.get("year")
                    ) for ed in data.get("education", []) if ed.get("degree")
                ]

                certifications = [
                    CertificationItem(name=c.get("name", "")) for c in data.get("certifications", []) if c.get("name")
                ]

                if projects:
                    return ExtractedResume(
                        name=data.get("name") or "Candidate",
                        email=data.get("email"),
                        phone=data.get("phone"),
                        location=data.get("location"),
                        summary=data.get("summary"),
                        skills=extracted_skills,
                        skill_categories=categories,
                        experience=experience,
                        projects=projects,
                        education=education,
                        certifications=certifications,
                        achievements=data.get("achievements", []),
                        raw_text=cleaned,
                        filename=filename
                    )
            except Exception as e:
                logger.warning(f"Gemini resume parsing fallback to advanced deterministic parser: {e}")

        # 2. Deterministic Fallback Parser
        contact = cls.extract_contact_info(cleaned)
        skills, categories = cls.extract_skills(cleaned)
        sections = cls.extract_sections(cleaned)
        
        projects = cls.parse_projects(sections.get("projects", ""), full_resume_text=cleaned)
        experience = cls.parse_experience(sections.get("experience", ""))
        education = cls.parse_education(sections.get("education", ""))
        
        # Certifications
        certs_text = sections.get("certifications", "")
        certifications = []
        if certs_text:
            for l in certs_text.split('\n'):
                t = l.strip().lstrip('•-* ')
                if len(t) > 4:
                    certifications.append(CertificationItem(name=t))

        # Achievements
        achieve_text = sections.get("achievements", "")
        achievements = [l.strip().lstrip('•-* ') for l in achieve_text.split('\n') if len(l.strip()) > 8]

        return ExtractedResume(
            name=contact["name"],
            email=contact["email"],
            phone=contact["phone"],
            location=contact["location"],
            summary=sections.get("summary", "")[:400] or None,
            skills=skills,
            skill_categories=categories,
            experience=experience,
            projects=projects,
            education=education,
            certifications=certifications,
            achievements=achievements,
            raw_text=cleaned,
            filename=filename
        )

    @staticmethod
    def calculate_score(resume: ExtractedResume, jd: Optional[Any] = None) -> ResumeScoreBreakdown:
        # Explainable score calculation
        skills_count = len(resume.skills)
        skills_score = min(100, max(20, skills_count * 8))

        proj_count = len(resume.projects)
        projects_score = min(100, max(20, proj_count * 25 + sum(len(p.technologies) for p in resume.projects) * 3))

        exp_count = len(resume.experience)
        experience_score = min(100, max(20, exp_count * 30 + sum(len(e.responsibilities) for e in resume.experience) * 4))

        edu_count = len(resume.education)
        education_score = 90 if edu_count > 0 else 50

        # Completeness based on contact info, sections presence
        completeness = 40
        if resume.email: completeness += 15
        if resume.phone: completeness += 15
        if resume.skills: completeness += 10
        if resume.projects: completeness += 10
        if resume.education: completeness += 10
        completeness_score = min(100, completeness)

        relevance_score = 80
        if jd and hasattr(jd, "required_skills") and jd.required_skills:
            jd_skills_lower = [s.lower() for s in jd.required_skills]
            match_count = sum(1 for s in resume.skills if s.lower() in jd_skills_lower)
            relevance_score = min(100, int((match_count / max(1, len(jd.required_skills))) * 100))

        overall = int(
            (skills_score * 0.25) +
            (projects_score * 0.25) +
            (experience_score * 0.20) +
            (education_score * 0.10) +
            (completeness_score * 0.10) +
            (relevance_score * 0.10)
        )

        strengths = []
        improvement_areas = []

        if skills_score >= 75:
            strengths.append(f"Broad technical skill set with {skills_count} verified tools & languages.")
        else:
            improvement_areas.append("Add more specific technical frameworks, libraries, and databases.")

        if proj_count >= 2:
            strengths.append(f"Demonstrated practical hands-on capability across {proj_count} detailed projects.")
        else:
            improvement_areas.append("Include at least 2-3 end-to-end projects showcasing architecture and deployment.")

        if exp_count > 0:
            strengths.append("Contains structured professional work experience with documented responsibilities.")
        else:
            improvement_areas.append("Highlight open-source contributions, internships, or freelance work.")

        if completeness_score >= 85:
            strengths.append("Complete resume structure with clear contact information and education.")
        
        rationale = (
            f"Your resume achieved an overall score of {overall}/100. "
            f"You have strong scores in {'Skills' if skills_score >= 70 else 'Education'} ({max(skills_score, education_score)}/100) "
            f"and {'Projects' if projects_score >= 70 else 'Experience'} ({max(projects_score, experience_score)}/100). "
            f"{'Matching target job profile.' if relevance_score >= 75 else 'Can be enhanced by tailoring project descriptions with target job keywords.'}"
        )

        return ResumeScoreBreakdown(
            overall_score=overall,
            skills_score=skills_score,
            projects_score=projects_score,
            experience_score=experience_score,
            education_score=education_score,
            completeness_score=completeness_score,
            relevance_score=relevance_score,
            strengths=strengths,
            improvement_areas=improvement_areas,
            rationale=rationale
        )
